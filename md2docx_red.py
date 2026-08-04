#!/usr/bin/env python3
# Convert a red-marked markdown source to .docx.
# Sentinels:  [[R ... R]] -> red text ;  **bold** -> bold ;  # / ## / ### headings ; |pipe tables|
import sys, re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

RED = RGBColor(0xC0, 0x00, 0x00)
SRC, OUT = sys.argv[1], sys.argv[2]
doc = Document()
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10.5)

def add_inline(par, text):
    # split into red / non-red on [[R ... R]]
    for i, chunk in enumerate(re.split(r'\[\[R(.*?)R\]\]', text, flags=re.S)):
        red = (i % 2 == 1)
        if chunk == '':
            continue
        # handle bold within chunk
        for j, seg in enumerate(re.split(r'\*\*(.*?)\*\*', chunk, flags=re.S)):
            if seg == '':
                continue
            r = par.add_run(seg)
            if red:
                r.font.color.rgb = RED
            if j % 2 == 1:
                r.bold = True

def flush_table(rows):
    if not rows:
        return
    ncol = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=ncol)
    t.style = 'Table Grid'
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci in range(ncol):
            txt = row[ci] if ci < len(row) else ''
            p = cells[ci].paragraphs[0]
            add_inline(p, txt.strip())
            for run in p.runs:
                run.font.size = Pt(8.5)
                if ri == 0:
                    run.bold = True

lines = open(SRC, encoding='utf-8').read().split('\n')
i = 0
tbuf = []
def is_table_line(s): return s.strip().startswith('|') and s.strip().endswith('|')
while i < len(lines):
    line = lines[i]
    if is_table_line(line):
        # collect table block
        block = []
        while i < len(lines) and is_table_line(lines[i]):
            block.append(lines[i]); i += 1
        rows = []
        for b in block:
            cells = [c for c in b.strip().strip('|').split('|')]
            if all(set(c.strip()) <= set('-: ') and c.strip() for c in cells):
                continue  # separator row
            rows.append(cells)
        flush_table(rows)
        continue
    s = line.rstrip()
    if s.startswith('### '):
        p = doc.add_paragraph(); r = p.add_run(s[4:]); r.bold = True; r.font.size = Pt(11.5)
    elif s.startswith('## '):
        p = doc.add_paragraph(); r = p.add_run(s[3:]); r.bold = True; r.font.size = Pt(13)
    elif s.startswith('# '):
        p = doc.add_paragraph(); r = p.add_run(s[2:]); r.bold = True; r.font.size = Pt(16)
    elif s.strip() == '':
        pass
    elif s.lstrip().startswith('- '):
        p = doc.add_paragraph(style='List Bullet'); add_inline(p, s.lstrip()[2:])
    else:
        p = doc.add_paragraph(); add_inline(p, s)
    i += 1

doc.save(OUT)
print("saved", OUT, "| paragraphs", len(doc.paragraphs), "| tables", len(doc.tables))
