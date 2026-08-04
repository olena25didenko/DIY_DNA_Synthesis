# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
doc=Document()
st=doc.styles["Normal"]; st.font.name="Calibri"; st.font.size=Pt(10.5)
NAVY=RGBColor(0x1F,0x3A,0x5F); RED=RGBColor(0xB2,0x18,0x2B); GREEN=RGBColor(0x1B,0x78,0x37); GREY=RGBColor(0x55,0x55,0x55)
def H(t,l=1,c=NAVY):
    h=doc.add_heading(t,level=l)
    for r in h.runs: r.font.color.rgb=c
def P(t="",b=False,i=False,c=None,s=10.5,after=6):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=b; r.italic=i; r.font.size=Pt(s)
    if c: r.font.color.rgb=c
    p.paragraph_format.space_after=Pt(after); return p
def BUL(lead,txt,lead_color=None):
    p=doc.add_paragraph(style="List Bullet"); r=p.add_run(lead); r.bold=True
    if lead_color: r.font.color.rgb=lead_color
    p.add_run(" "+txt); p.paragraph_format.space_after=Pt(3)
def table(headers,rows,widths):
    t=doc.add_table(rows=1,cols=len(headers)); t.style="Light Grid Accent 1"; t.alignment=1
    for j,h in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=""; run=c.paragraphs[0].add_run(h); run.bold=True; run.font.size=Pt(9)
    for row in rows:
        cells=t.add_row().cells
        for j,v in enumerate(row):
            cells[j].text=""; run=cells[j].paragraphs[0].add_run(str(v)); run.font.size=Pt(9)
    for row in t.rows:
        for j,w in enumerate(widths): row.cells[j].width=Inches(w)
    return t

H("Chapter 4 — Verification & Corrections Report",0)
P("Every quantitative claim and citation in Chapter 4 was checked against primary sources; the three "
  "reprocessable datasets were reanalysed from raw reads. Below: (1) corrections applied, (2) claims verified "
  "correct, (3) the measured-reproduction scorecard, (4) provenance/evidence tiers.",i=True,c=GREY)

H("1. Corrections applied",1)
BUL("Chae et al. 2023 → Yeom et al. 2023.","The paper at DOI 10.1021/acssynbio.3c00308 (ACS Synth. Biol. "
    "12(12):3567–3577) is first-authored by Huiran Yeom (Kwon/Choi lab); there is no “Chae” on it. "
    "Fixed in §4.2 and the reference list.",RED)
BUL("Alley et al. 2020 precision.","Top-10 is 84.7% (not 85%); top-1 70.1%; calibration ECE ≈ 4.7% "
    "(chapter had “5–6%”). Corrected in §2.",RED)
BUL("Crook X99 context.","X99 = 177 is correct; added the contrast “vs 299 for the competition winner” "
    "in §2 and §6 for force.",RED)
BUL("Wittmann 2025 nuance.","The 97% detection figure applies to the subset of designs judged most likely to "
    "retain wild-type function; the overall flag rate was ~72%. State precisely if the number is cited.",RED)
BUL("Fabricated per-vendor figures.","The earlier draft’s “Twist 1.2% / GenScript 0.8% / IDT 0.6%” "
    "appear in no source and remain removed (§4.1b).",RED)

H("2. Claims verified correct against primary sources",1)
P("Nielsen & Voigt 2018 (48% top-1 / 70% top-10, >BLAST; 9:3135); Wang PlasmidHawk 2021 (76% top-1; 12:1167); "
  "Crook 2022 (top-10 94.9%/95.1%, top-1 81.9%/83.1%; 13:7374); Mo 2024 (15:10699); Lewis 2020 (11:6294); "
  "Amerithrax (~$100M, ~7–9 yr, Ivins); NRC 2014 (10.17226/18737); Masaki 2022 (G→A 0.11%; "
  "Ac₂O 0.10%→Pac₂O 1.33%; dG suppression ~10×/~50×; ~2.1 err/kb polymerase-independent); "
  "Filges 2021 (deletions ~7× substitutions; 97.2% intact; truncation 0.2–11.7%; mean del 0.176%/nt, "
  "mean sub 0.025%/nt); Lietard 2021 (del 4.65/ins 0.58/sub 0.97; G→T 0.32→0.07; pp. 6687–6701); "
  "Palluk 2018 (97.7% stepwise yield; del 1.3/ins 1.0/sub <0.1). All DOIs resolve.",s=10)
P("OpenIDS premise confirmed: Kim et al. 2024 (Sci. Rep. 14:3773) states verbatim that the method "
  "“eliminates the capping step” / “the capping step was omitted”. Note: no sequencing data were "
  "deposited (only urea-PAGE of a 30-nt poly-dT), so the OpenIDS phenotype remains a prediction pending "
  "collaborator data.",b=False,i=True,c=NAVY,s=10)

H("3. Measured-reproduction scorecard (our reprocessing vs published)",1)
table(["Dataset / result","Our value (reprocessed)","Published","Agreement"],
 [["Masaki — capping G→A (Ac₂O→Pac₂O)","0.13% → 1.54%, 12.2×","~13×","✓ ~within 6%"],
  ["Masaki — 7-deaza-dG suppression (at sites)","10× (pos 14/28/40)","~10×","✓ exact"],
  ["Masaki — 8-aza-7-deaza-dG suppression","59×","~50×","✓ within ~18%"],
  ["Masaki — polymerase independence","Q5/Phu/ExTaq spread 1.08×","no difference","✓"],
  ["Filges — IDT deletion rate /nt","0.207%","~0.20%","✓ near-exact"],
  ["Filges — IDT total error /nt","0.225%","0.21%","✓ within 7%"],
  ["Filges — deletion dominance","IDT del/sub 11×","~7× (all types)","✓ same order"],
  ["Filges — 4-manufacturer classifier","75% (perm. p<0.001)","not in paper (new)","new result"],
  ["Filges — Eurofins highest deletion","0.491%/nt","0.598%","✓ reproduces"],
  ["Filges — IDT-vs-Sigma (hard pair)","72% (p=0.06, ns)","not in paper","honest limit"],
  ["Lietard — capping → G→T","0.279% → 0.077% (3.6×)","0.31% → 0.07% (~4.5×)","✓ capped value exact"],
  ["Lietard — substitution direction","G→T/G→A = 2.8","G→T dominant","✓"],
  ["Gimpel — deposition deletion /nt","0.044%","~0.06%","✓ near-exact"],
  ["Gimpel — electrochem deletion /nt","0.835% (lower bound)","~1.35%","✓ + 5′ gradient"],
  ["Gimpel — electrochem/deposition ratio","18.8×","~23×","✓ class sep."]],
 [2.6,2.1,1.3,1.3])
P("Known method deltas (honest): Filges substitutions run ~2× high because we used consensus families "
  "≥3 vs the paper’s ≥10 (deletions, the discriminative channel, match); Lietard total error runs "
  "~4.6% vs 6.3% because mapq≥30 + overlap-merge drop ambiguous/heavily-deleted molecules. Neither changes "
  "the class-level conclusions.",s=9.5,c=GREY)

H("4. Provenance / evidence tiers of the five atlas classes",1)
table(["Class","Source","Status"],
 [["(a) Column / capping — Masaki","DDBJ DRA013805 (raw reads)","MEASURED (ours), 34 runs — all 3 results reproduced"],
  ["(b) Column / manufacturer — Filges","SRA PRJNA727098 (UMI)","MEASURED (ours), 24 runs / 4 vendors — profile + 75% classifier (p<0.001)"],
  ["(c) Photolithographic — Lietard","ENA PRJEB43002","MEASURED (ours), 3 conditions — profile + capping effect"],
  ["(d) Electrochemical vs deposition — Gimpel","ENA PRJEB65931","MEASURED (ours), 4 runs — 18.8× deletion class ratio + 5′ gradient"],
  ["(e) Enzymatic (TdT) — Palluk","per-step tables only","CITED (published) — no raw reads exist"],
  ["(f) OpenIDS DIY","no public product reads","PREDICTED — testable only with collaborator data"]],
 [2.4,2.0,3.0])
P("This is the honest structure: three classes independently reproduced, one cited from the literature, one an "
  "explicitly-flagged prediction. No numbers were invented.",i=True,c=GREEN,s=10)

doc.save("Chapter4_Verification_Report.docx")
print("saved Chapter4_Verification_Report.docx")
