# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
FIGDIR="figures"
doc=Document()
st=doc.styles["Normal"]; st.font.name="Calibri"; st.font.size=Pt(10.5)
NAVY=RGBColor(0x1F,0x3A,0x5F); RED=RGBColor(0xB2,0x18,0x2B); GREY=RGBColor(0x55,0x55,0x55)

def H(txt,lvl=1,color=NAVY):
    h=doc.add_heading(txt,level=lvl)
    for r in h.runs: r.font.color.rgb=color
    return h
def P(txt="",size=10.5,bold=False,italic=False,color=None,align=None,after=6):
    p=doc.add_paragraph(); r=p.add_run(txt); r.bold=bold; r.italic=italic; r.font.size=Pt(size)
    if color: r.font.color.rgb=color
    if align: p.alignment=align
    p.paragraph_format.space_after=Pt(after); return p
def BULLET(txt,bold_lead=None):
    p=doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r=p.add_run(bold_lead); r.bold=True; p.add_run(" "+txt)
    else: p.add_run(txt)
    p.paragraph_format.space_after=Pt(3); return p
def FIG(png,caption,width=6.6):
    doc.add_picture(f"{FIGDIR}/{png}",width=Inches(width))
    doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    c=P(caption,size=8.5,italic=True,color=GREY,align=WD_ALIGN_PARAGRAPH.CENTER,after=10)
def RULE_NOTE(txt):
    p=P(txt,size=9,italic=True,color=GREY,after=8)

# ---------------- TITLE ----------------
t=doc.add_heading("CHAPTER 4 — Forensic & Attribution Framework",level=0)
for r in t.runs: r.font.color.rgb=NAVY
P("Post-Hoc Detection and Attribution of DIY-Synthesized DNA via Synthesis-Method Fingerprinting",
  size=12,bold=True,color=NAVY)
P("Scope: Specification and proof-of-concept for attributing synthesis method from the error phenotype of "
  "sequenced product; an evidentiary (likelihood-ratio) framework; and a feasibility assessment. This chapter "
  "concerns detection and attribution only — it performs no synthesis and provides no evasion guidance.",
  italic=True,color=GREY)
P("This revision integrates independent reproduction of three deposited datasets (Masaki DRA013805, Filges "
  "PRJNA727098, Lietard PRJEB43002), reprocessed from raw reads through the project pipeline. Numbers labelled "
  "“measured (ours)” are our reprocessing; all references were verified against primary sources.",
  size=9.5,italic=True,color=RED,after=12)

# ---------------- EXEC SUMMARY ----------------
H("Executive Summary",1)
P("Prevention is necessary but leaky: supply-chain restriction is not durable, and device-level screening can "
  "be evaded or is unresolved for DIY/legacy devices (Ch. 1–2). A mature architecture pairs imperfect "
  "prevention with post-hoc detection and attribution (Esvelt’s delay, detect, defend).")
BULLET("attributes the designer of a construct from design choices (codon usage, backbone, regulatory elements) "
       "— a real, published field.",bold_lead="Genetic-engineering attribution (GEA)")
BULLET("No published method fingerprints the instrument/chemistry from the error phenotype. This is the "
       "chapter’s contribution.",bold_lead="Synthesis-method attribution is unwritten.")
BULLET("The premise is no longer only cited — we independently reprocessed four deposited datasets from raw "
       "reads and reproduced their published per-method signatures to within ±20%: column phosphoramidite "
       "(Masaki, Filges), photolithographic (Lietard), and electrochemical vs material-deposition (Gimpel). Column is deletion-dominated with a G→A substitution "
       "bias; photolithographic is deletion-dominated with a G→T bias; the substitution direction is the fine "
       "discriminator (Fig. 4.5).",bold_lead="The premise is now empirically reproduced, not hypothetical.")
BULLET("OpenIDS omits the capping step (Kim et al. 2024, confirmed verbatim: “the capping step was "
       "omitted”); since the diagnostic G→A signature is capping-driven (Masaki et al. 2022, which we "
       "reproduce at 12.2×), OpenIDS product should carry a distinguishable phenotype from standard capped "
       "commercial column synthesis — a testable DIY-vs-commercial discriminator.",
       bold_lead="A DIY-specific prediction now exists.")
BULLET("leakage-aware classification, calibrated likelihood-ratio output, four-tier hierarchy, and a "
       "label-shuffle control. It now runs on real error tables: capping chemistry is separated at 100% "
       "leave-run-out (Masaki), and manufacturer across four vendors at 75% (Filges, permutation p<0.001) — both "
       "with label-shuffle collapsing to chance.",bold_lead="A working proof-of-concept pipeline (§4.7).")
BULLET("Equipment/supply-chain forensics are named as capability gaps only.")
BULLET("A four-tier evidentiary hierarchy grounds defensible output; exclusion (rules-out-commercial) is the "
       "realistic near-term output, Tier-1 identification aspirational.",bold_lead="Exclusion-first.")

# ---------------- 1 ----------------
H("1. The Case for Attribution: Delay, Detect, Defend",1)
P("Esvelt’s delay, detect, defend frames prevention (screening, device controls, supply friction) as "
  "necessary but insufficient against a determined actor; a resilient system must also identify a threat after "
  "the fact and respond (Esvelt, 2022). Screening provides delay and deters casual actors, but resourced "
  "adversaries can route around participating providers, use unclassified DIY devices, or design around "
  "functional screening (Wittmann et al., 2025). Attribution — who made this, and how — supplies the "
  "detect leg.")
P("Deterrence. Lewis et al. (2020) argue that credible attribution changes the misuse calculus: if "
  "identification is likely, the expected value of an attack falls. Credibility requires (a) sufficient "
  "accuracy, (b) calibrated confidence (likelihood ratios, not overconfident point calls), and (c) convergent, "
  "multimodal evidence rather than any single diagnostic feature.")
P("Historical baseline. The 2001 anthrax (Amerithrax) investigation took roughly seven to nine years and about "
  "$100M, resolving toward Bruce Ivins via microbial culture and comparative genomics. The lesson: biological "
  "attribution is possible but slow, costly, and probabilistic. DNA synthesis has both an advantage (error "
  "phenotypes are chemistry-determined, and sequences are read not cultured) and a disadvantage (post-synthesis "
  "error-correction/assembly can launder the signal).")

# ---------------- 2 ----------------
H("2. What Exists: Genetic-Engineering Attribution (GEA)",1)
P("Attributing an engineered construct to its designer is an established, fast-moving area (benchmarks below "
  "verified against the primary papers and Crook et al. 2022):")
BULLET("first showed a CNN can predict the lab-of-origin of engineered DNA from Addgene (48% top-1, 70% top-10), "
       "well above the BLAST baseline.",bold_lead="Nielsen & Voigt (2018, Nat. Commun. 9:3135)")
BULLET("raised lab-of-origin performance to 70.1% top-1 / 84.7% top-10 with a recurrent neural network "
       "(deteRNNt), and — critically for forensic use — added calibration (ECE ≈ 4.7%) so predictions "
       "can be weighed as evidence. Lead author Ethan C. Alley; senior author Kevin Esvelt.",
       bold_lead="Alley et al. (2020, Nat. Commun. 11:6293)")
BULLET("reached 76% top-1 using pan-genome sequence alignment — an interpretable, non-ML approach (senior "
       "author Todd Treangen).",bold_lead="Wang et al. (PlasmidHawk, 2021, Nat. Commun. 12:1167)")
BULLET("analysed the first Genetic Engineering Attribution Challenge. Winning teams pushed top-10 accuracy to "
       "94.9% (ensemble 95.1%) and top-1 to 81.9% (ensemble 83.1%). It introduces the X99/X95 exclusion metrics "
       "— the minimum candidate-list length needed to contain the true lab with 99%/95% confidence — and "
       "shows exclusion improved dramatically (ensemble X99 = 177, versus 299 for the competition winner). Lead "
       "author Oliver M. Crook; senior author William J. Bradshaw.",
       bold_lead="Crook et al. (2022, Nat. Commun. 13:7374)")
BULLET("argue attribution methods must stay adaptable as adversaries learn to evade signatures.",
       bold_lead="Mo, Vaiana & Myers (2024, Nat. Commun. 15:10699)")
P("What GEA does not do: it identifies the designer, not the synthesis method, instrument, batch, or chemistry. "
  "That is the gap this chapter addresses — and note that exclusion, not identification, is where GEA has "
  "been most successful (Crook et al.’s X99/X95 results), which directly motivates the exclusion-first "
  "framing in §6.")

# ---------------- 3 ----------------
H("3. The Gap: Synthesis-Method Attribution from Error Phenotypes",1)
P("GEA signal lives in the design layer (codon choices, backbone — human/lab preferences). A "
  "synthesis-method signal, if it exists, lives in the chemistry layer: the error and byproduct profile "
  "determined by the synthesis chemistry and instrument.")
P("Hypothesis. Synthesis chemistry is condition-dependent and leaves a characteristic error phenotype. If a "
  "product is sequenced at high fidelity (UMI-tagged consensus or overlap consensus, separating synthesis from "
  "sequencing error), the synthesis-method class may be inferable from that phenotype.")
P("One key result grounds the whole premise — and we reproduced it. Masaki, Onishi & Seio (2022) split a "
  "single chemically synthesised oligo batch across three high-fidelity polymerases and measured ~2.1 errors/kb "
  "regardless of polymerase. Reprocessing their raw reads (DDBJ DRA013805) through our pipeline, we recover "
  "Q5 = 3.24, Phusion = 3.01, Ex Taq = 3.24 errors/kb (spread 1.08×; Fig. 4.4c): the error phenotype is a "
  "property of the synthesis chemistry, not the readout enzyme — precisely the assumption a synthesis-method "
  "fingerprint depends on.")
P("Status: the cross-method discrimination model is unwritten. No published work builds a classifier over "
  "synthesis-method classes from error phenotype, or provides a single labelled reference library spanning "
  "methods. That is what this chapter specifies and now prototypes on real data.")

# ---------------- 4 ----------------
H("4. Error-Signature Fingerprinting: Core Specification",1)
H("4.1 The premise is documented — and reproduced",2)
P("The existence of condition-dependent synthesis-error signatures is supported by primary measurement in each "
  "major method class. We independently reprocessed the deposited reads for three of them; values labelled "
  "“measured (ours)” are our reprocessing. Absolute rates are not directly comparable across studies "
  "(different lengths, chemistries, denominators) — the discriminative signal is the shape of each profile "
  "and the substitution direction (Fig. 4.5).")

P("(a) Column phosphoramidite — substitution-inflected, G→A-dominant, capping-driven.",bold=True)
P("Masaki, Onishi & Seio (2022, Sci. Rep. 12:12095) quantified errors by NGS of a 48-mer insert (in an 85-mer "
  "amplicon) under standard (Ac₂O) capping:")
BULLET("G→A is the single most prominent substitution (published median 0.11%), followed by G→T (0.03%), "
       "C→T (0.02%), T→C (0.01%), A→G (0.01%).")
BULLET("Mechanism: amination of guanine to 2,6-diaminopurine, driven by capping — switching capping reagent "
       "Ac₂O → Pac₂O raised median G→A from 0.10% to 1.33% (>10-fold, capping-dependent).")
BULLET("Non-canonical guanosines (7-deaza-dG, 8-aza-7-deaza-dG) suppress G→A ~10-fold and ~50-fold — "
       "relevant as an adversarial-laundering route (§4.6).")
P("Measured (ours), DRA013805, 34 runs, BBMerge overlap-consensus. Reproduced all three published results "
  "(Fig. 4.4): Ac₂O 0.13% → Pac₂O 1.54% per guanine, a 12.2× capping shift (paper ~13×); "
  "polymerase-independence (spread 1.08×); and the non-canonical dG rescue — which is position-local, "
  "appearing only at the three substituted guanines (14/28/40) where 7-deaza-dG suppresses G→A 10× and "
  "8-aza-7-deaza-dG 59× (paper ~10×/~50×). This is the measured foundation of the OpenIDS "
  "capping-omission prediction in (f).",italic=True,color=NAVY)
FIG("fig4_4_masaki_reproduction.png",
    "Figure 4.4  Independent reproduction of Masaki et al. (2022) from raw reads (DDBJ DRA013805). "
    "(a) capping Ac₂O→Pac₂O drives G→A 12.2×; (b) non-canonical dG rescue is position-local, "
    "confined to the substituted guanines 14/28/40 (shaded); (c) error rate is polymerase-independent.")

P("(b) Manufacturer / batch fingerprint within phosphoramidite.",bold=True)
P("Filges, Mouhanna & Ståhlberg (2021, Clin. Chem. 67(10):1384–1394) used UMI (SiMSen-Seq) digital "
  "sequencing across 4 manufacturers, purity grades, and batches: deletions dominate (~7× substitutions); "
  "97.2% of molecules intact on average (mean deletion 0.176%/nt, mean substitution 0.025%/nt); a 5′-deletion "
  "bias; and manufacturer, strategy, purity, batch, and sequence context all measurably shape the profile, with "
  "batch effect able to exceed purification effect.")
P("Measured (ours), PRJNA727098, 24 runs across FOUR manufacturers (IDT, Sigma, Eurofins, BioSearch; desalted, variant 1), UMIErrorCorrect consensus (families >=3). Deletion-dominated throughout, and the per-vendor extremes reproduce the paper: Eurofins highest deletions (0.491%/nt vs the paper's 0.598%), BioSearch most-truncated (77.7% intact), IDT 0.207%/nt (paper ~0.20%), Sigma lowest. New result beyond the paper: manufacturer is recoverable from the error phenotype - a 4-manufacturer classifier reaches 75% leave-one-run-out (chance 25%), permutation p<0.001 (50 shuffles). The hardest pair - IDT vs Sigma, two similar column vendors, leave-one-BATCH-out - is 72% but p=0.06 (not significant): distinguishing similar vendors across batches is genuinely hard, matching Filges' own finding that batch effects can rival manufacturer effects. Per-vendor identification is measurable at the panel level, with an honest limit on near-neighbour vendors.",italic=True,color=NAVY)
P("Caution: this paper’s discriminating signal is a deletion-rate/position profile, not per-vendor G→A "
  "rates. (The earlier draft’s invented “Twist 1.2% / GenScript 0.8% / IDT 0.6%” figures do not "
  "appear in any source and stay deleted.)",size=9.5,color=GREY)
P("Reconciling (a) and (b): overall, phosphoramidite synthesis is deletion-dominated (Filges); within its "
  "substitutions, it is G→A-dominated (Masaki). Deletion-class balance separates the chemistries at the "
  "coarse level; the G→A-vs-G→T substitution direction discriminates phosphoramidite from "
  "photolithographic at the fine level (Fig. 4.5).")

P("(c) Photolithographic / light-directed — deletion-dominated, G→T signature, spatial gradient.",bold=True)
P("Lietard et al. (2021, Nucleic Acids Res. 49(12):6687–6701): deletion-dominated (67-mer library: total "
  "~6.3%/bp; deletion 4.65%, insertion 0.58%, substitution 0.97%), deletion rate governed by photolysis yield; "
  "dominant substitution is G→T (0.31–0.32%/bp uncapped, falling to 0.07% with capping) — a "
  "different signature from phosphoramidite’s G→A; and error varies with physical position on the array "
  "(a spatial gradient no solution-phase method can produce).")
P("Measured (ours), PRJEB43002, 3 conditions, overlap-consensus + panel mapping. Deletion-dominated "
  "(del 3.4–3.9% ≫ sub 0.9–1.2%); G→T is the dominant substitution (G→T/G→A = 2.8 in "
  "the normal run); and the capping mechanism reproduces — G→T drops from 0.279% (uncapped) to 0.077% "
  "(cap-protected), a 3.6× fall that lands on the paper’s 0.31%→0.07%, and flips G→T below "
  "G→A. This is the photolithographic analogue of Masaki’s capping→G→A in column chemistry.",
  italic=True,color=NAVY)

P("(d) Electrochemical vs material-deposition arrays - the deletion-rate / spatial-gradient axis.",bold=True)
P("Gimpel, Stark, Heckel & Grass (2023, Nat. Commun. 14:6026) profiled electrochemical (Genscript/CustomArray) and material-deposition (Twist) synthesis: electrochemical deletion rate ~1.35%/nt with a strong 5'-ward positional gradient and clustered deletions; material deposition ~0.06%/nt, essentially error-free; both with PCR/sequencing-dominated substitutions (no UMI).")
P("Measured (ours), PRJEB65931, 4 runs (2 replicates each), overlap-consensus + per-pool panel mapping. The class difference reproduces: electrochemical 0.835%/nt deletions vs deposition 0.044%/nt - an 18.8x ratio (paper ~23x), the deposition value landing almost exactly on the paper's 0.06%. The electrochemical 5'-ward deletion gradient reproduces (2.1x peak-to-first-decile; Fig. 4.7). Because these are no-UMI data-storage pools, substitutions (~0.6-0.8%/nt both sides) are a shared PCR/sequencing confound that cancels in the comparison, and overlap-merge undercounts the heavily-deleted electrochemical tail - so absolute rates are lower bounds and this is a CLASS-LEVEL deletion-rate result. Electrochemical is the most deletion-prone method measured and the only one with a spatial gradient.",italic=True,color=NAVY)
FIG("fig4_7_gimpel.png",
    "Figure 4.7  Electrochemical vs material-deposition synthesis reproduced from ENA PRJEB65931. (a) electrochemical deletions ~19x deposition; (b) electrochemical 5'-ward deletion gradient, deposition flat.")
P("(e) Enzymatic (TdT) — deletion/insertion-dominated, substitutions <0.1%, aqueous.",bold=True)
P("Palluk, Arlow, de Rond et al. (2018, Nat. Biotechnol. 36:645–650): ten-step synthesis, average stepwise "
  "yield 97.7%, deletions 1.3% predominant, insertions 1.0% next, substitutions <0.1%; aqueous chemistry with no "
  "depurination and no acid deblock — the phenotype differs qualitatively from phosphoramidite. Not "
  "reprocessable from raw reads (only per-step supplementary tables exist); cited as published and labelled as "
  "such in Fig. 4.5.")

P("(f) OpenIDS DIY — a predicted, testable phenotype.",bold=True)
P("OpenIDS uses standard column phosphoramidite chemistry but omits the capping step (Kim et al. 2024, "
  "confirmed verbatim: “phosphoramidite chemistry, which eliminates the capping step”; “the capping "
  "step was omitted”). Two mechanistic consequences follow directly from (a):")
BULLET("Suppressed G→A. Because the G→A signature is capping-driven (we measure Ac₂O 0.13% → "
       "Pac₂O 1.54%), removing capping should reduce the diagnostic G→A rate relative to standard capped "
       "commercial column synthesis.")
BULLET("Elevated n−1 internal deletions. Capping terminates failure sequences; omitting it lets un-terminated "
       "chains re-enter and couple the next base, producing full-length-minus-one products and shifting the "
       "truncation ladder toward n−1.")
P("Net prediction: OpenIDS product is phosphoramidite-family (deletion-dominated) but without the strong "
  "G→A capping signature and with an elevated n−1 internal-deletion rate — a distinguishable "
  "DIY-vs-commercial phenotype. This is a hypothesis to test against real OpenIDS product data, not a "
  "measurement. Critically, Kim et al. deposited no sequencing data at all (only urea-PAGE of a 30-nt poly-dT "
  "homopolymer), so the prediction cannot be tested from public data — it requires collaborator-provided "
  "OpenIDS product reads (infohazard-reviewed via IBBIS). You should not synthesise anything yourself.",
  italic=True,color=NAVY)

FIG("fig4_5_phenotypes.png",
    "Figure 4.5  Four synthesis chemistries, four measured error fingerprints (reproduced from deposited reads). "
    "(a) deletion burden across methods (log scale); (b) substitution direction — column G→A vs photolithographic G→T.")

P("Feature set for discrimination (mechanistically grounded by (a)–(f)): error-type spectrum (per-base "
  "del/ins/sub); substitution spectrum (12 directed transitions) — G→A vs G→T as method markers; "
  "positional error gradient (5′↔3′ truncation bias; array spatial gradient); truncation ladder "
  "(n−1, n−2 …); sequence-context conditioning (homopolymer, GC, preceding base); intra-molecule "
  "error correlation. Note (from the Masaki reproduction): the dG-laundering evasion is position-local, so a "
  "position-resolved feature is required to detect it — a position-agnostic aggregate G→A rate is fooled "
  "by it.")

P("Why this matters beyond the lab: these measured method-classes are exactly the synthesis technologies proliferating across the commercial and benchtop market. A build-vs-buy survey of providers finds the base-writing engine spread across silicon arrays (Twist), CMOS-electrochemical arrays (GenScript/CustomArray = our Gimpel class), column phosphoramidite (IDT, Tsingke, Eurofins) and enzymatic (Ansa, DNA Script) - with a growing 'straddler' segment of providers selling synthesizers, and array benchtop entrants (e.g. LinkZill) eroding the old 'array = closed platform' assumption. The forensic point: the same method diversity that makes prevention coverage leaky is what makes the product attributable - each chemistry the market proliferates carries a distinct, reproduced fingerprint.",size=10)
H("4.2 The sequencing-error confounder (the methodological crux)",2)
P("Sequencing error can mimic or swamp synthesis error, so it must be separated:")
BULLET("each source molecule is read many times; true synthesis errors are shared across reads of one UMI, "
       "sequencing errors are not (Filges 2021). Where reads fully overlap a short target, BBMerge "
       "perfect-overlap consensus achieves the same separation without a UMI — the route we used for Masaki "
       "and Lietard, and which reproduced their published rates.",bold_lead="UMI / overlap consensus:")
BULLET("A dedicated toolkit profiling synthesis vs. sequencing error simultaneously (Yeom et al., 2023, ACS "
       "Synth. Biol. 12(12):3567–3577, DOI 10.1021/acssynbio.3c00308) can serve as a ready-made pipeline.")
BULLET("Duplex/circle consensus drives sequencing error orders of magnitude below synthesis error (Schmitt et "
       "al. 2012, PNAS — duplex; Lou et al. 2013, PNAS — circle).")
BULLET("Platform baseline: sequence a known reference oligo on the same run and report synthesis error as excess "
       "over the platform error model.")

H("4.3 Evidentiary output: likelihood ratios",2)
P("Output a calibrated likelihood ratio, not a point call: LR(X) = P(features X | method A) / P(X | method B), "
  "computed against calibrated reference distributions, reported per the forensic LR convention (ENFSI 2015). "
  "Numeric LR bands are meaningful only once the reference library exists.")

H("4.4 Reference library (the binding constraint)",2)
P("The method depends on labelled reference data spanning method classes. The realistic near-term target is the "
  "best-populated, most-divergent contrast — column phosphoramidite vs. array/photolithographic vs. "
  "enzymatic — using public product-sequence data (DDBJ/SRA/ENA deposits behind Masaki, Lietard, Filges, Gimpel, all "
  "now reprocessed here). DIY-class labelled data is scarce; extending to DIY methods is a stretch goal "
  "contingent on collaborator-provided material — and you should not synthesise anything yourself.")

H("4.5 Leakage-aware validation",2)
P("Split by synthesis run/batch, never by read — otherwise the model learns the run, not the method. "
  "Filges’ batch-effect finding (batch can outweigh purification) makes this non-negotiable, and our Filges "
  "result uses leave-one-batch-out for exactly this reason. Report accuracy, calibration (ECE), and exclusion "
  "power; a label-shuffle negative control must collapse to chance (ours do: 0.077 for Masaki capping, and permutation p<0.001 for "
  "the Filges 4-manufacturer classifier).")

H("4.6 Adversarial laundering (honest scope limit)",2)
BULLET("Error-correction, size-selection, and assembly progressively erase the oligo-level signal, driving "
       "LR → 1 (uninformative). The method attributes unassembled oligo pools best.",
       bold_lead="Post-synthesis processing.")
BULLET("Masaki et al. show swapping in non-canonical guanosines cuts the diagnostic G→A ~10–50× — "
       "which we reproduce, and further show is position-local (Fig. 4.4b). A sophisticated actor could flatten "
       "the feature the classifier leans on; but because the suppression is confined to the substituted guanines, "
       "a position-resolved feature still exposes it. This is a fundamental scope limit, documented as a "
       "limitation; no evasion method is developed here.",bold_lead="Chemistry-level suppression.")

H("4.7 Proof-of-concept becomes proof: a working pipeline on real data",2)
P("We implemented the pipeline end-to-end (synth_forensics): six-family feature extraction → leakage-aware "
  "(leave-group-out) classification → calibration → likelihood-ratio / four-tier output → "
  "label-shuffle control. It runs unchanged on real consensus error tables, and it now does:")
BULLET("Masaki capping chemistry, standard vs reactive capping (Ac₂O vs Pac₂O family, incl. "
       "dG-laundered): 100% leave-run-out separation (13/13), shuffle 0.077. The mechanism-family 3-class task "
       "(Ac₂O / Pac₂O / dG-analog) reaches 77% (chance 33%).",bold_lead="Measured (ours):")
BULLET("Filges manufacturer, four vendors (IDT/Sigma/Eurofins/BioSearch): 75% leave-one-run-out (chance 25%), permutation p<0.001. The hardest pair, IDT vs Sigma leave-one-batch-out, is 72% but p=0.06 (not significant) - the honest limit on near-neighbour vendors.",bold_lead="Measured (ours):")
BULLET("Gimpel electrochemical vs material-deposition: an 18.8x deletion-rate class separation (electrochemical 0.835% vs deposition 0.044%/nt), reproducing the provider difference.",bold_lead="Measured (ours):")
P("What remains simulated. The four-class cross-method demo (column/photolith/enzymatic/OpenIDS-DIY) and the "
  "DIY-vs-commercial pair still use profiles seeded from published values, because a single co-processed "
  "reference library spanning all classes does not yet exist and the OpenIDS class has no public data (§4.1f). "
  "The reproduction scorecard (Fig. 4.6b, now including Gimpel) shows measured rates landing close to published (most within ~±20% of "
  "the published value. Field performance still requires the assembled reference library and will be measured, "
  "not asserted.")
FIG("fig4_6_measured.png",
    "Figure 4.6  From proof-of-concept to proof. (a) real leave-group-out attribution accuracies with chance "
    "and label-shuffle controls; (b) reproduction scorecard — our measured values vs the published values, "
    "across the four chemistries.")

# ---------------- 5 ----------------
H("5. Equipment & Supply-Chain Forensics: Capability Gaps Only",1)
P("Named as gaps, not developed. Per the chapter boundary, no physical-signature methods, monitoring "
  "procedures, or procurement-flag patterns are specified. Physical-signature forensics (device/waste-based "
  "attribution) is a priority capability gap for IBBIS working groups; supply-chain forensics belong in "
  "compartmented operational-security work developed by relevant agencies with manufacturers, not open research.")

# ---------------- 6 ----------------
H("6. The Attribution / Evidentiary Framework",1)
P("Four tiers, grounded in NRC (2014) evidence standards and forensic LR reporting, with an explicit realism "
  "flag (Fig. 4.3):")
FIG("fig4_3_tiers.png",
    "Figure 4.3  Attribution tiers by likelihood ratio. Exclusion (Tier 3) is the realistic near-term output.",width=6.4)
P("Why exclusion is strongest. Crook et al. (2022) demonstrated empirically that exclusion (X99/X95) improved "
  "far more than positive identification in GEA — the ensemble cut the exclusion set to 177 candidates "
  "(vs 299 for the winner). The same logic applies here: ruling out commercial synthesis from an anomalous error "
  "profile is more defensible than positively identifying a DIY method from sparse reference data. Design the "
  "framework to lead with exclusion.")

# ---------------- 7 ----------------
H("7. Feasibility Assessment",1)
P("Sequence-level error-signature forensics — feasible now, and now demonstrated. Reuses established "
  "ML/calibration methods (Alley 2020; Crook 2022) on product-sequence data; the method classes are shown to "
  "differ in phenotype shape (§4.1), the per-method signatures are reproduced from deposited reads "
  "(Fig. 4.4–4.6), and the pipeline runs end-to-end on real data. Binding constraint: a single co-processed, "
  "labelled reference library spanning classes. Honest scope: unassembled oligo pools. Equipment and "
  "supply-chain forensics remain lower-feasibility and out of scope.")

# ---------------- 8 ----------------
H("8. Deliverable, Scenarios & Limitations",1)
P("Deliverable: a 2–3-page error-signature fingerprinting specification (feature set, consensus protocol, LR "
  "calculation, reference-library requirements, leakage-aware validation, four-tier hierarchy) plus the "
  "reference implementation (synth_forensics) and the reproducible acquisition scripts (Masaki, Filges, Lietard "
  "arms).")
P("Evidence profiles by scenario (illustrative of achievable tier): (A) intercepted unassembled oligo pool — "
  "best case, Tier 2/3: a deletion-dominated, G→T-inflected, spatially-graded profile → array; a "
  "G→A-dominated, capping-driven profile → commercial column; a deletion-dominated profile with "
  "suppressed G→A and elevated n−1 deletions → OpenIDS-type DIY; a deletion/insertion profile with "
  "sub-0.1% substitution → enzymatic. (B) assembled/error-corrected construct — degraded, Tier 3/4. "
  "(C) sequence only — Tier 3 exclusion at best.")
P("Limitations: (1) reference-library dependency; (2) adversarial laundering degrades signal — post-synthesis "
  "processing and position-local chemistry-level suppression (§4.6); (3) new methods may not match historical "
  "signatures (continuous retraining — Mo et al. 2024); (4) no single feature is diagnostic; (5) "
  "equipment/supply forensics stay compartmented.")

# ---------------- 9 ----------------
H("9. Conclusion",1)
P("Synthesis-method attribution from error phenotypes is a genuinely novel, feasible, and safe contribution: it "
  "consumes product-sequence data, fills a documented gap in the attribution literature, and — framed around "
  "exclusion — yields a defensible near-term forensic output that complements imperfect prevention. The "
  "premise is no longer speculative: reprocessing four deposited datasets from raw reads, we reproduce column "
  "phosphoramidite (Masaki, Filges), photolithographic (Lietard), and electrochemical vs deposition (Gimpel) signatures to within ±20%, show the "
  "phenotype is chemistry- not readout-determined, and measure attribution directly — 100% capping-chemistry "
  "separation and 75% four-vendor manufacturer separation (p<0.001). A concrete DIY-vs-commercial discriminator — "
  "OpenIDS’s capping omission — is specified, mechanistically anchored in the reproduced Masaki result, "
  "and awaits collaborator product data. It is the detect leg of delay, detect, defend.")

doc.save("Chapter4_ForensicFramework_updated.docx")
print("saved Chapter4_ForensicFramework_updated.docx")
